import json
from django.contrib import admin
from django.contrib.auth.admin import UserAdmin
from django.contrib.auth.models import User
from .models import (
    University, College, Subject, Question, Option, PDFResource,
    TestSession, StudentAnswer, UserProfile, FavoriteQuestion,
    ActivationCode, ExamDataJSONUpload,
    Deck, Flashcard, UserFlashcardProgress, FlashcardDocxUpload,
    CollegeStudyGuide, CollegeStudyGuideUpload,
)

# ... (the rest of the classes stay untouched, I need to append at the bottom)

@admin.register(University)
class UniversityAdmin(admin.ModelAdmin):
    list_display = ('name', 'code', 'is_active')
    list_filter = ('is_active',)
    search_fields = ('name', 'code')
    list_editable = ('is_active',)

@admin.register(Subject)
class SubjectAdmin(admin.ModelAdmin):
    list_display = ('name', 'code', 'is_active')
    list_filter = ('is_active',)
    search_fields = ('name', 'code')
    list_editable = ('is_active',)

@admin.register(College)
class CollegeAdmin(admin.ModelAdmin):
    list_display = ('name', 'university', 'code', 'category', 'icon_name', 'high_school_weight', 'exam_weight', 'is_active')
    list_filter = ('university', 'is_active', 'category')
    search_fields = ('name', 'university__name', 'code')
    list_editable = ('high_school_weight', 'exam_weight', 'is_active')
    filter_horizontal = ('subjects',)

class OptionInline(admin.TabularInline):
    model = Option
    extra = 4

@admin.register(Question)
class QuestionAdmin(admin.ModelAdmin):
    list_display = ('get_short_text', 'subject', 'grade_level', 'unit', 'exam_year', 'source_university', 'times_appeared')
    list_filter = ('subject', 'grade_level', 'unit', 'exam_year', 'source_university')
    search_fields = ('text', 'explanation')
    inlines = [OptionInline]
    list_per_page = 50

    def get_short_text(self, obj):
        return str(obj)[:50] + '...' if len(str(obj)) > 50 else str(obj)
    get_short_text.short_description = 'السؤال'

@admin.register(PDFResource)
class PDFResourceAdmin(admin.ModelAdmin):
    list_display = ('title', 'college', 'subject', 'resource_type', 'year')
    list_filter = ('college', 'resource_type', 'year', 'subject')
    search_fields = ('title', 'college__name')

class StudentAnswerInline(admin.TabularInline):
    model = StudentAnswer
    extra = 0
    readonly_fields = ('question', 'selected_option', 'is_correct')
    can_delete = False
    def has_add_permission(self, request, obj=None):
        return False

@admin.register(TestSession)
class TestSessionAdmin(admin.ModelAdmin):
    list_display = ('student', 'test_type', 'exam_mode', 'college', 'score_percentage', 'created_at')
    list_filter = ('test_type', 'exam_mode', 'college', 'created_at')
    search_fields = ('student__username', 'college__name')
    readonly_fields = ('student', 'test_type', 'exam_mode', 'college', 'total_questions', 'correct_answers_count', 'score_percentage')
    inlines = [StudentAnswerInline]
    def has_add_permission(self, request):
        return False

@admin.register(FavoriteQuestion)
class FavoriteQuestionAdmin(admin.ModelAdmin):
    list_display = ('user', 'question', 'created_at')
    list_filter = ('user', 'question__subject')
    search_fields = ('user__username', 'question__text')
    readonly_fields = ('user', 'question', 'created_at')


@admin.register(ActivationCode)
class ActivationCodeAdmin(admin.ModelAdmin):
    list_display = ('id', 'code_display_safe', 'is_used', 'used_by', 'used_at', 'duration_days', 'note', 'created_at')
    list_filter = ('is_used', 'duration_days')
    search_fields = ('note', 'used_by__username', 'code_display')
    readonly_fields = ('code_hash', 'is_used', 'used_by', 'used_at', 'created_at')
    actions = ['generate_5_codes', 'generate_10_codes']

    def code_display_safe(self, obj):
        """يعرض الكود فقط إذا لم يُستخدم بعد"""
        if obj.code_display and not obj.is_used:
            return f"\u2b50 {obj.code_display}"
        return "\u2014 (مُشفّر)"
    code_display_safe.short_description = 'الكود'

    def generate_5_codes(self, request, queryset):
        codes = []
        for _ in range(5):
            obj, plain = ActivationCode.generate(admin_name=request.user.username, note='تم توليده من لوحة التحكم')
            codes.append(plain)
        self.message_user(request, f"تم توليد 5 أكواد: {', '.join(codes)}")
    generate_5_codes.short_description = '➕ توليد 5 أكواد جديدة'

    def generate_10_codes(self, request, queryset):
        codes = []
        for _ in range(10):
            obj, plain = ActivationCode.generate(admin_name=request.user.username, note='تم توليده من لوحة التحكم')
            codes.append(plain)
        self.message_user(request, f"تم توليد 10 أكواد: {', '.join(codes)}")
    generate_10_codes.short_description = '➕ توليد 10 أكواد جديدة'


# إدارة ملفات الطلاب
class UserProfileInline(admin.StackedInline):
    model = UserProfile
    can_delete = False
    verbose_name_plural = 'ملف اشتراك الطالب'

class CustomUserAdmin(UserAdmin):
    inlines = (UserProfileInline, )

admin.site.unregister(User)
admin.site.register(User, CustomUserAdmin)
admin.site.register(UserProfile)

@admin.register(ExamDataJSONUpload)
class ExamDataJSONUploadAdmin(admin.ModelAdmin):
    list_display = ('id', 'uploaded_at', 'processed')
    readonly_fields = ('processed', 'log')

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)
        if not obj.processed and obj.file:
            try:
                obj.file.open('rb')
                file_content = obj.file.read().decode('utf-8-sig')
                data = json.loads(file_content)
                obj.file.close()

                meta = data.get('metadata', {})
                university_code = meta.get('university_code')
                college_code = meta.get('college_code')
                subject_code = meta.get('subject_code')
                exam_year = meta.get('exam_year')

                # التحقق من الأكواد
                university = University.objects.filter(code=university_code).first()
                if not university:
                    raise Exception(f"لم يتم العثور على الجامعة بالرمز: {university_code}")

                college = College.objects.filter(code=college_code, university=university).first()
                if not college:
                    raise Exception(f"لم يتم العثور على الكلية بالرمز: {college_code} لجامعة {university_code}")

                subject = Subject.objects.filter(code=subject_code).first()
                if not subject:
                    raise Exception(f"لم يتم العثور على المادة بالرمز: {subject_code}")

                # التأكد من ارتباط المادة بالكلية 
                if not college.subjects.filter(id=subject.id).exists():
                    college.subjects.add(subject)

                questions_data = data.get('questions', [])
                added_count = 0

                for q_data in questions_data:
                    # تلافي التكرار باستخدام النص والسنة والمادة
                    q, created = Question.objects.get_or_create(
                        text=q_data.get('text', '').strip(),
                        subject=subject,
                        exam_year=exam_year,
                        source_university=university,
                        source_college=college,
                        defaults={
                            'grade_level': q_data.get('grade_level', 12),
                            'unit': q_data.get('unit', 1),
                            'explanation': q_data.get('explanation', ''),
                        }
                    )

                    if created:
                        added_count += 1
                        correct_identifier = q_data.get('correct_identifier', '')
                        options_list = q_data.get('options', [])
                        
                        for opt in options_list:
                            Option.objects.create(
                                question=q,
                                identifier=opt.get('identifier', '').strip().upper(),
                                text=opt.get('text', '').strip(),
                                is_correct=(opt.get('identifier', '').strip().upper() == correct_identifier.strip().upper())
                            )

                obj.processed = True
                obj.log = f"نجاح: تم ربط {added_count} سؤال لـ {college.name} - مادة {subject.name} - سنة {exam_year}."
                obj.save()
            except Exception as e:
                obj.log = f"خطأ أثناء المعالجة: {str(e)}"
                obj.save()


# --------------------------------------------------------
# نظام البطاقات التعليمية (Flashcards Admin)
# --------------------------------------------------------

class FlashcardInline(admin.TabularInline):
    model = Flashcard
    extra = 3
    fields = ('front', 'back', 'order')


@admin.register(Deck)
class DeckAdmin(admin.ModelAdmin):
    list_display = ('name', 'subject', 'card_count', 'created_at')
    list_filter = ('subject',)
    search_fields = ('name', 'subject__name')
    inlines = [FlashcardInline]

    def card_count(self, obj):
        return obj.cards.count()
    card_count.short_description = 'عدد البطاقات'


@admin.register(FlashcardDocxUpload)
class FlashcardDocxUploadAdmin(admin.ModelAdmin):
    list_display = ('deck_name', 'subject', 'uploaded_at', 'processed')
    readonly_fields = ('processed', 'log')

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)
        if not obj.processed and obj.file:
            try:
                from docx import Document
                import io

                obj.file.open('rb')
                file_bytes = obj.file.read()
                obj.file.close()

                doc = Document(io.BytesIO(file_bytes))

                # ── استخراج الصور وحفظها ──
                import os
                import uuid
                from django.conf import settings
                
                # إنشاء مجلد الصور إذا لم يكن موجوداً
                flashcard_media_path = os.path.join(settings.MEDIA_ROOT, 'flashcards')
                if not os.path.exists(flashcard_media_path):
                    os.makedirs(flashcard_media_path)

                def get_image_url(rId):
                    try:
                        image_part = doc.part.related_parts[rId]
                        image_bytes = image_part.blob
                        ext = image_part.content_type.split('/')[-1]
                        filename = f"fc_{uuid.uuid4().hex}.{ext}"
                        with open(os.path.join(flashcard_media_path, filename), 'wb') as f:
                            f.write(image_bytes)
                        return f"{settings.MEDIA_URL}flashcards/{filename}"
                    except:
                        return None

                raw_lines = []
                for para in doc.paragraphs:
                    line_parts = []
                    for run in para.runs:
                        img_url = None
                        if run._element is not None:
                            # البحث عن blip (الذي يحتوي على rId الصورة) بصورة دقيقة جداً
                            blips = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                            for blip in blips:
                                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if rId:
                                    # محاولة استخراج الصورة الفعلية
                                    extracted_url = get_image_url(rId)
                                    if extracted_url:
                                        img_url = extracted_url
                                        break # اكتفينا بصورة واحدة لهذا الـ run

                        if img_url:
                            line_parts.append(f'[IMG:{img_url}]')
                        elif run.text:
                            # إضافة النص فقط إذا لم تكن هناك صورة ناجحة
                            line_parts.append(run.text)

                    full_line = ''.join(line_parts).strip()
                    if full_line:
                        raw_lines.append(full_line)

                # ── تحليل Front:/Back: مع دعم العناوين ──
                cards_data = []
                current_front = None
                current_back = None
                current_context = ""
                reading = None  # 'front' أو 'back' أو None

                for line in raw_lines:
                    stripped = line.strip()
                    
                    if not stripped:
                        if reading == 'front' and current_front is not None:
                            current_front += '\n'
                        elif reading == 'back' and current_back is not None:
                            current_back += '\n'
                        continue

                    # فحص إذا كان السطر عبارة عن فاصل (مثل ___ أو ---)
                    is_separator = len(stripped) >= 3 and set(stripped).issubset({'-', '_', '=', '*'})
                    if is_separator:
                        reading = None  # إنهاء قراءة البطاقة الحالية، ما يلي سيكون عنواناً أو بطاقة جديدة
                        continue

                    lower_stripped = stripped.lower()
                    if lower_stripped.startswith('front:') or lower_stripped.startswith('front :'):
                        # حفظ البطاقة السابقة
                        if current_front is not None and current_back is not None:
                            cards_data.append((current_front.strip(), current_back.strip()))
                        
                        # بداية بطاقة جديدة
                        remainder = stripped.split(':', 1)[1].strip() if ':' in stripped else ''
                        if current_context:
                            current_front = f"[{current_context}]\n{remainder}"
                        else:
                            current_front = remainder
                            
                        current_back = None
                        reading = 'front'
                        
                    elif lower_stripped.startswith('back:') or lower_stripped.startswith('back :'):
                        remainder = stripped.split(':', 1)[1].strip() if ':' in stripped else ''
                        current_back = remainder
                        reading = 'back'
                        
                    else:
                        # سطر عادي
                        if reading == 'front' and current_front is not None:
                            current_front += '\n' + stripped
                        elif reading == 'back' and current_back is not None:
                            current_back += '\n' + stripped
                        else:
                            # نحن لسنا بصدد قراءة سؤال أو إجابة (إما في بداية الملف أو بعد فاصل)
                            # إذن هذا السطر هو عنوان فرعي (Context) للبطاقات القادمة
                            current_context = stripped

                # حفظ البطاقة الأخيرة
                if current_front is not None and current_back is not None:
                    cards_data.append((current_front.strip(), current_back.strip()))

                if not cards_data:
                    raise Exception("لم يتم العثور على أي بطاقات. تأكد من صيغة Front:/Back: في الملف.")

                # ── إنشاء Deck والبطاقات ──
                deck, _ = Deck.objects.get_or_create(
                    name=obj.deck_name.strip(),
                    subject=obj.subject,
                )

                added = 0
                for i, (front, back) in enumerate(cards_data):
                    if front and back:
                        Flashcard.objects.get_or_create(
                            deck=deck,
                            front=front,
                            defaults={'back': back, 'order': i + 1}
                        )
                        added += 1

                obj.processed = True
                obj.log = f"✅ نجاح: تم إضافة {added} بطاقة إلى مجموعة «{deck.name}» — مادة {obj.subject.name}"
                obj.save()

            except Exception as e:
                obj.log = f"❌ خطأ: {str(e)}"
                obj.save()


# --------------------------------------------------------
# دليل المراجعة (Study Guide Admin)
# --------------------------------------------------------

@admin.register(CollegeStudyGuide)
class CollegeStudyGuideAdmin(admin.ModelAdmin):
    list_display = ('title', 'college', 'subject', 'is_active', 'updated_at')
    list_filter = ('college', 'subject', 'is_active')
    search_fields = ('title', 'content')
    list_editable = ('is_active',)


@admin.register(CollegeStudyGuideUpload)
class CollegeStudyGuideUploadAdmin(admin.ModelAdmin):
    list_display = ('guide_title', 'college', 'subject', 'uploaded_at', 'processed')
    readonly_fields = ('processed', 'log')

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)
        if not obj.processed and obj.file:
            try:
                from docx import Document
                import io

                obj.file.open('rb')
                file_bytes = obj.file.read()
                obj.file.close()

                doc = Document(io.BytesIO(file_bytes))

                # استخراج النص من الملف مع الحفاظ على البنية
                lines = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        lines.append(text)

                if not lines:
                    raise Exception("الملف فارغ أو لا يحتوي على نصوص.")

                content = '\n'.join(lines)

                # إنشاء أو تحديث الدليل
                guide, created = CollegeStudyGuide.objects.update_or_create(
                    college=obj.college,
                    subject=obj.subject,
                    defaults={
                        'title': obj.guide_title.strip(),
                        'content': content,
                        'is_active': True,
                    }
                )

                action = "إنشاء" if created else "تحديث"
                obj.processed = True
                obj.log = f"✅ تم {action} دليل «{guide.title}» بنجاح — {len(lines)} سطر."
                obj.save()

            except Exception as e:
                obj.log = f"❌ خطأ: {str(e)}"
                obj.save()
